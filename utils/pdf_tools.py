import time

# from gpts.gpt_agent import profileAgent

import io
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_formatted_text(paragraph, text):
    """
    Add text to a paragraph, parsing **bold** markdown syntax.
    
    Example: "This is **bold text** and normal" 
    -> Creates runs with appropriate formatting
    """
    # Split by **bold** pattern while keeping the delimiters
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Bold text - remove the ** markers
            bold_text = part[2:-2]
            run = paragraph.add_run(bold_text)
            run.bold = True
        elif part:  # Skip empty strings
            # Normal text
            paragraph.add_run(part)


def markdown_table_to_docx(markdown_text: str, output_path: str = None, logo_path: str = None):
    """
    Convert markdown to Docx with a logo positioned at top-left.
    - Left: -3cm indent
    - Height: Moved up by reducing header distance to 0.5cm
    - Supports **bold** markdown syntax in regular text

    If output_path is provided, saves to disk.
    Returns a BytesIO object containing the document.
    """

    lines = markdown_text.strip().split('\n')
    doc = Document()
    
    # --- LOGO POSITIONING START ---
    if logo_path:
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        
        # 1. VERTICAL POSITION (Height)
        # "Header distance" is the gap from the top edge of the paper to the start of the header.
        # Default is usually ~1.27cm. Setting it to 0.5cm moves the logo UP.
        section.header_distance = Cm(0.5)
        
        # 2. HORIZONTAL POSITION (Left)
        # Align left and use negative indent to pull it into the margin.
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header_para.paragraph_format.left_indent = Cm(-3)
        
        # Remove any extra spacing that might push it down
        header_para.paragraph_format.space_before = Pt(0)
        header_para.paragraph_format.space_after = Pt(0)
        
        # 3. INSERT IMAGE
        run = header_para.add_run()
        try:
            # Adjust width as needed
            run.add_picture(logo_path, width=Inches(1))
        except FileNotFoundError:
            print(f"Warning: Logo file not found at {logo_path}")
    # --- LOGO POSITIONING END ---
    
    # Process text content
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Detect table start
        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                line = lines[i].strip()
                if '---' not in line:
                    table_lines.append(line)
                i += 1
            
            if table_lines:
                rows = []
                for line in table_lines:
                    cells = [c.strip() for c in line.strip('|').split('|')]
                    rows.append(cells)
                
                if rows:
                    max_cols = max(len(row) for row in rows)
                    
                    # Pad rows
                    for row in rows:
                        while len(row) < max_cols:
                            row.append('')
                    
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                cell = table.rows[row_idx].cells[col_idx]
                                # Clear default paragraph and add formatted text
                                cell.text = ''
                                cell_para = cell.paragraphs[0]
                                add_formatted_text(cell_para, cell_data)
                                
                                # Make header row bold
                                if row_idx == 0:
                                    for run in cell_para.runs:
                                        run.bold = True
                    doc.add_paragraph()
        
        # Headings
        elif line in [
            '1. Business Overview', '2. Key Stakeholders', '3. Revenue Split', 
            '4a. Products/Services Overview', '4b. Geographical Footprint',
            '5. Key Developments', '5. Key Recent Developments','6. Financial Highlights', '7. Capital Structure']:
            para = doc.add_paragraph()
            add_formatted_text(para, line)
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(16)
            i += 1

        elif line in ['Summary / Interpretation', 'Summary', 'Sources:', 'Sources']:
            para = doc.add_paragraph()
            add_formatted_text(para, line)
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(12)
            i += 1
        
        # Bullet points
        elif line.startswith('-') or line.startswith('•'):
            para = doc.add_paragraph(style='List Bullet')
            add_formatted_text(para, line[1:].strip())
            i += 1
        
        # Regular paragraphs
        elif line:
            para = doc.add_paragraph()
            add_formatted_text(para, line)
            i += 1
        else:
            i += 1

    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Optionally save to disk if output_path is provided
    if output_path:
        doc.save(output_path)
        print(f"✓ Saved to: {output_path}")

    return buffer
